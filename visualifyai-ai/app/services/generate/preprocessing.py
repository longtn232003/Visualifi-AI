from app.const import MODEL_NAME_TOKENIZER
from app.schemas.response import ResponseAPI
from app.const import StatusResponse

from transformers import AutoTokenizer
import pdfplumber
import docx

class PreprocessingService:
    def __init__(self):
        self.tokenizer = AutoTokenizer.from_pretrained(MODEL_NAME_TOKENIZER)

    async def extract_text_from_pdf(self, file):
        with pdfplumber.open(file.file) as pdf_file:
            # get text from pdf
            full_text = ""
            for page in pdf_file.pages:
                if page.images:
                    results = ResponseAPI(
                        StatusResponse.HTTP_BAD_REQUEST,
                        {"message": "PDF contains images"},
                    )
                    return results
                if page.find_tables():
                    results = ResponseAPI(
                        StatusResponse.HTTP_BAD_REQUEST,
                        {"message": "PDF contains tables"},
                    )
                    return results

                text = page.extract_text()
                full_text = full_text + text

            return full_text
        
    async def extract_text_from_docx(self, file):
        if not hasattr(file.file, "seekable"):
            file.file.seekable = lambda: True

        file.file.seek(0)
        docx_file = docx.Document(file.file)
        full_text = ""
        # check image in docx
        if docx_file.inline_shapes:
            results = ResponseAPI(
                StatusResponse.HTTP_BAD_REQUEST,
                {"message": "DOCX contains images"},
            )
            return results
        
        # check table in docx
        if docx_file.tables:
            results = ResponseAPI(
                StatusResponse.HTTP_BAD_REQUEST,
                {"message": "DOCX contains tables"},
            )
            return results
        
        for paragraph in docx_file.paragraphs:
            full_text = full_text + paragraph.text

        return full_text
        
    async def extract_text_from_txt(self, file):    
        with open(file.file, "r", encoding="utf-8") as txt_file:
            full_text = txt_file.read()
            return full_text
        
    async def token_length(self, text):
        token_ids = self.tokenizer.encode(text, padding=True, truncation=True, return_tensors='pt')
        return len(token_ids[0])